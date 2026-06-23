"""Extract text and structural metadata from CV files (PDF, DOCX, TXT)."""
from __future__ import annotations

import io
from dataclasses import dataclass


@dataclass
class ParsedCV:
    text: str
    # Structural signals an ATS parser cares about.
    page_count: int = 0
    word_count: int = 0
    has_images: bool = False
    has_tables: bool = False
    parse_error: str = ""


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def parse_cv(filename: str, data: bytes) -> ParsedCV:
    """Dispatch on file extension. Never raises — errors land in parse_error."""
    ext = _ext(filename)
    try:
        if ext == "pdf":
            return _parse_pdf(data)
        if ext == "docx":
            return _parse_docx(data)
        if ext in ("txt", "text", "md"):
            return _parse_txt(data)
        return ParsedCV(text="", parse_error=f"Unsupported file type: .{ext}")
    except Exception as exc:  # noqa: BLE001 — surface, don't crash the batch
        return ParsedCV(text="", parse_error=f"{type(exc).__name__}: {exc}")


def _finalize(text: str, **kwargs) -> ParsedCV:
    text = text.strip()
    return ParsedCV(text=text, word_count=len(text.split()), **kwargs)


def _parse_pdf(data: bytes) -> ParsedCV:
    import pdfplumber

    parts: list[str] = []
    has_images = False
    has_tables = False
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            if page.images:
                has_images = True
            if page.find_tables():
                has_tables = True
    return _finalize(
        "\n".join(parts),
        page_count=page_count,
        has_images=has_images,
        has_tables=has_tables,
    )


def _parse_docx(data: bytes) -> ParsedCV:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    has_tables = len(doc.tables) > 0
    # Pull text out of tables too (recruiters often put skills in a table).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    # python-docx exposes inline shapes for embedded images.
    has_images = len(doc.inline_shapes) > 0
    return _finalize(
        "\n".join(paragraphs),
        page_count=1,
        has_images=has_images,
        has_tables=has_tables,
    )


def _parse_txt(data: bytes) -> ParsedCV:
    text = data.decode("utf-8", errors="replace")
    return _finalize(text, page_count=1)
